import re, sys
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from difflib import SequenceMatcher

def read_docx(doc_path):
    """
    Read the content of a Word document and return it as a string.

    :param doc_path: The path of the Word document to read.
    :return: The content of the document as a string.
    """
    # Create a Document object from the file
    doc = Document(doc_path)
    
    # Initialize an empty string to store the text
    text = ''
    
    # Iterate over all paragraphs in the document
    for paragraph in doc.paragraphs:
        # Concatenate the text of each paragraph separated by a newline character
        text += paragraph.text + '\n'
    
    # Return the text
    return text

def compare_docs(doc_path1, doc_path2):
    """
    Compare two Word documents and return the character positions where they differ.

    :param doc_path1: The path of the first Word document to compare.
    :param doc_path2: The path of the second Word document to compare.
    :return: A list of tuples representing the start and end positions of differences between the documents.
    """
    # Read the content of the documents
    text1 = read_docx(doc_path1)
    text2 = read_docx(doc_path2)
    
    # Split the texts into words and whitespace using regular expressions
    words1 = re.findall(r'\S+|\s+', text1)
    words2 = re.findall(r'\S+|\s+', text2)

    # Create a SequenceMatcher object to compare the sequences of words
    s = SequenceMatcher(None, words1, words2)

    # Initialize an empty list to store the positions of differences
    diff_positions = []

    # Initialize a variable to keep track of the current character position
    char_pos = 0
    
    # Iterate over the operations needed to transform one sequence into the other
    for tag, i1, i2, j1, j2 in s.get_opcodes():
        # If the operation is not 'equal'
        if tag != 'equal':
            # The start position of the difference is the current character position
            start = char_pos

            # The end position is calculated by adding the length of the different words
            end = start + len(''.join(words2[j1:j2]))

            # Add a tuple with the start and end positions to the list of differences
            diff_positions.append((start, end))

        # Update the current character position by adding the length of the current words
        char_pos += len(''.join(words2[j1:j2]))
    
    # Return the list of positions where differences were found
    return diff_positions

def get_tuples_in_run(run_start, run_end, diff_positions):
    """
    This function returns the given tuples whose positions are totally
    or partially in the current run.
    """
    # Defining list of tuples in current run
    new_diff_positions = []

    # Iterating through all tuples
    for position_start, position_end in diff_positions:
        # Checking if current tuple is past current run
        # As they are sorted, if current tuple is past current run,
        # all of the following ones will also be
        if run_end < position_start:
            break

        # If start position of tuple is inside current run, tuple is at least
        # partially contained in the current run
        if run_start <= position_start < run_end:
            if run_start <= position_end <= run_end:
                # If end position is inside current run, it is totally contained in the run
                new_diff_positions.append((position_start, position_end))
            else:
                # If end position is not inside current run, it is partially contained in the run
                # As they are sorted, if the end of a tuple is outside the run, the rest of them
                # will also be
                new_diff_positions.append((position_start, run_end))
                break
        else:
            # If start position of tuple is before current run but end position is inside,
            # it is partially contained in the run
            if position_start < run_start and position_end <= run_end:
                new_diff_positions.append((run_start, position_end))

    # Returning result positions
    return new_diff_positions

def get_current_positions(diff_positions, run_end):
    """
    This function removes all the tuples already past the last run.
    """
    # Defining new list of positions
    new_positions = []

    # Iterating through list of positions
    for position_start, position_end in diff_positions:
        if run_end >= position_end:
            # If start of tuple was before or inside last run, it is no longer needed
            continue
        else:
            # If start of tuple is past last run it is still needed
            new_positions.append((position_start, position_end))
    
    return new_positions

def add_run_with_style(paragraph, text, run, highlight=False):
    """
    This function creates a run with the given text with style based on
    the given run, and optionally highlighting in yellow.
    """
    # Applying styles
    new_run = paragraph.add_run(text)
    new_run.bold = run.bold
    new_run.element = run.element

    # Font styles
    new_run.font.all_caps = run.font.all_caps
    new_run.font.bold = run.font.bold
    new_run.font.complex_script = run.font.complex_script
    new_run.font.cs_bold = run.font.cs_bold
    new_run.font.cs_italic = run.font.cs_italic
    new_run.font.double_strike = run.font.double_strike
    new_run.font.emboss = run.font.emboss
    new_run.font.hidden = run.font.hidden
    new_run.font.highlight_color = run.font.highlight_color
    new_run.font.imprint = run.font.imprint
    new_run.font.italic = run.font.italic
    new_run.font.math = run.font.math
    new_run.font.name = run.font.name
    new_run.font.no_proof = run.font.no_proof
    new_run.font.outline = run.font.outline
    new_run.font.rtl = run.font.rtl
    new_run.font.shadow = run.font.shadow
    new_run.font.size = run.font.size
    new_run.font.small_caps = run.font.small_caps
    new_run.font.snap_to_grid = run.font.snap_to_grid
    new_run.font.spec_vanish = run.font.spec_vanish
    new_run.font.strike = run.font.strike
    new_run.font.subscript = run.font.subscript
    new_run.font.superscript = run.font.superscript
    new_run.font.underline = run.font.underline
    new_run.font.web_hidden = run.font.web_hidden

    # Font color styles
    new_run.font.color.rgb = run.font.color.rgb
    new_run.font.color.theme_color = run.font.color.theme_color

    new_run.italic = run.italic
    new_run.style = run.style
    new_run.underline = run.underline
    # Conditionally highlighting
    if highlight:
        new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

def add_new_runs(paragraph, reference_run, tuples_in_run, run_start):
    """
    This function recieves a paragraph, a reference run, and a run text,
    together with a list of tuples contained in that run and the start position
    of that run, and it adds the necessary runs to include the tuples with special
    style.
    """
    # Getting reference text
    text_to_spread = reference_run.text

    # Getting the tuples's position relative to the run
    relative_tuples = []
    for tuple_start, tuple_end in tuples_in_run:
        relative_tuples.append((tuple_start - run_start, tuple_end - run_start))
    
    # Splitting the text into separate strings
    split_text = []
    last_pos = 0

    # Iterating throug the relative tuples
    for start_index, end_index in relative_tuples:
        # If the last index is smaller than the start of the current tuple,
        # insert the string from the last index to the start of the tuple
        if start_index > last_pos:
            split_text.append((text_to_spread[last_pos:start_index], False))
        
        # Insert the string of the tuple
        split_text.append((text_to_spread[start_index:end_index], True))

        # Updating index
        last_pos = end_index
    
    # If after finishing with all the tuples, there is some of the string left,
    # add it at the end
    if last_pos < len(text_to_spread):
        split_text.append((text_to_spread[last_pos:], False))
    
    # Adding runs to the document
    for text, highlight in split_text:
        add_run_with_style(paragraph, text, reference_run, highlight=highlight)

def create_highlighted_doc(doc_path, diff_positions, output_filename):
    """
    This function receives a reference docx file, an output file name,
    and a list of tuples containing the start and end of each string to be
    highlighted in the text of the document.
    """
    # Creating reference and writable documents
    referenceDoc = Document(doc_path)
    doc2 = Document()

    # Defining starting positoon of the first run
    run_start = 0

    # Iterating through each paragraph of the reference document
    for paragraph in referenceDoc.paragraphs:
        # Creating a new paragraph and applying style in the new document
        # based on the reference one
        new_paragraph = doc2.add_paragraph()
        new_paragraph.alignment = paragraph.alignment

        # Iterating through each run in the current paragraph
        for run in paragraph.runs:
            # Calculating the absolute end position of the current run
            run_end = run_start + len(run.text)

            # Getting the list of tuples totally or partially contained in the current run
            # Only applies if the list is not empty
            if diff_positions:
                tuples_in_run = get_tuples_in_run(run_start, run_end, diff_positions)

            # Checking if there are any tuples contained in the current run 
            if tuples_in_run and diff_positions:
                add_new_runs(new_paragraph, run, tuples_in_run, run_start)
            else:
                # Adding new run with the same style as current run without highlight
                add_run_with_style(new_paragraph, run.text, run)

            # Updating list of tuples removing the ones already used
            # Only applies if the list is not empty
            if diff_positions:
                diff_positions = get_current_positions(diff_positions, run_end)

            # Updating absolute start position for current tun
            run_start = run_end
        
        # After a paragraph change, a \n is added to the text, which is not counted in the
        # runs text, so it needs to be accounted for outside
        run_start += 1
    
    # Saving the output in the defined document
    doc2.save(output_filename)

# Main function
if __name__ == '__main__':
    # Checking for --help option
    if len(sys.argv) == 2 and sys.argv[1] == '--help':
        print('Compares the text in file1.docx and file2.docx and creates a new Word document with the differences highlighted.')
        print('Specifically, the output will be a copy of the second input file, but it will have highlighted all the elements '
              'present in that file that are not present in the first file.')
        print('\tUsage: python3 docx_highlight_diff.py file1.docx file2.docx output_file.docx')
        sys.exit(0)

    # Checking parameter count
    if len(sys.argv) != 4:
        print('Error: Exactly three arguments must be provided. Use --help for more information.')
        sys.exit(1)
    
    # Making sure all introduced files are a .docx file
    for arg in [sys.argv[1], sys.argv[2]]:
        if not arg.endswith('.docx'):
            print(f'Error: {arg} is not a Word file. The extension must be .docx. Use --help for more information.')
            sys.exit(1)
    
    diff_texts = compare_docs(sys.argv[1], sys.argv[2])
    create_highlighted_doc(sys.argv[2], diff_texts, sys.argv[3])