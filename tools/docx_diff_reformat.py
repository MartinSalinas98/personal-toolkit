import sys
from docx import Document
from docx.text.run import Run
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn, nsmap
from docx.shared import RGBColor

# Global variables
# Tags used to differentiate between regular text and inserted/deleted text
INSERTED_OPENING_TAG = '<inserted>'
INSERTED_CLOSING_TAG = '</inserted>'
DELETED_OPENING_TAG = '<deleted>'
DELETED_CLOSING_TAG = '</deleted>'

# Modification type
MODIFICATION_INSERTED = 1
MODIFICATION_DELETED = 2

# Red RGB color
RED_COLOR = RGBColor(191, 4, 4)

def full_text(run):
    """
    This function receives a run in lxml.etree.Element format and returns the full text of the run,
    including inserted and/or deleted text.
    """
    text = ''

    # Add the mc namespace to nsmap
    nsmap['mc'] = 'http://schemas.openxmlformats.org/markup-compatibility/2006'

    # Check if the run is within an inserted or deleted element
    isInserted = run.getparent().tag == qn('w:ins')
    isDeleted = run.getparent().tag == qn('w:del')

    # If the run is within an inserted element, add an opening <inserted> tag to the output
    if isInserted:
        text += INSERTED_OPENING_TAG
    # If the run is within a deleted element, add an opening <deleted> tag to the output
    elif isDeleted:
        text += DELETED_OPENING_TAG

    # Iterate through the children of the run
    for child in run:
        # If the child is a text or delText element, add its text to the output
        if child.tag in (qn('w:t'), qn('w:delText')):
            text += child.text or ''
        # If the child is a tab element, add a tab character to the output
        elif child.tag == qn('w:tab'):
            text += '\t'
        # If the child is a line break or carriage return element, add a newline character to the output
        elif child.tag in (qn('w:br'), qn('w:cr')):
            text += '\n'
        # If the child is an AlternateContent element, recursively extract text from its first Choice element's paragraphs
        elif child.tag == qn('mc:AlternateContent'):
            for nestedP in child.xpath('mc:Choice[1]//w:p', namespaces=nsmap):
                text += full_text(nestedP)
                text += '\n'

    # If the run is within an inserted element, add a closing </inserted> tag to the output
    if isInserted:
        text += INSERTED_CLOSING_TAG
    # If the run is within a deleted element, add a closing </deleted> tag to the output
    elif isDeleted:
        text += DELETED_CLOSING_TAG

    return text

def add_run_with_style(paragraph, text, run, highlight=False, cross=False, color=None):
    """
    This function creates a run with the given text with style based on
    the given run, and optionally highlighting in yellow.
    """
    # Applying styles
    new_run = paragraph.add_run(text)
    new_run.bold = run.bold
    new_run.element = run.element
    new_run.italic = run.italic
    new_run.style = run.style
    new_run.underline = run.underline

    # Font styles
    new_run.font.all_caps = run.font.all_caps
    new_run.font.bold = run.font.bold
    new_run.font.complex_script = run.font.complex_script
    new_run.font.cs_bold = run.font.cs_bold
    new_run.font.cs_italic = run.font.cs_italic
    new_run.font.double_strike = run.font.double_strike
    new_run.font.emboss = run.font.emboss
    new_run.font.hidden = run.font.hidden
    new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW if highlight else run.font.highlight_color
    new_run.font.imprint = run.font.imprint
    new_run.font.italic = run.font.italic
    new_run.font.math = run.font.math
    new_run.font.name = run.font.name
    new_run.font.no_proof = run.font.no_proof
    new_run.font.outline = run.font.outline
    new_run.font.rtl = run.font.rtl
    new_run.font.shadow = run.font.shadow
    new_run.font.size = run.font.size
    new_run.font.small_caps = run.font.small_caps
    new_run.font.snap_to_grid = run.font.snap_to_grid
    new_run.font.spec_vanish = run.font.spec_vanish
    new_run.font.strike = True if cross else run.font.strike
    new_run.font.subscript = run.font.subscript
    new_run.font.superscript = run.font.superscript
    new_run.font.underline = run.font.underline
    new_run.font.web_hidden = run.font.web_hidden

    # Font color styles
    new_run.font.color.rgb = color if color else run.font.color.rgb
    new_run.font.color.theme_color = run.font.color.theme_color

def get_text_with_type(text):
    """
    This function receives a text that can be unmodified, inserted, or deleted,
    and returns that text with the modification it was applied to it.
    """
    # Detecting inserted text
    if text.startswith(INSERTED_OPENING_TAG) and text.endswith(INSERTED_CLOSING_TAG):
        return text[len(INSERTED_OPENING_TAG):-len(INSERTED_CLOSING_TAG)], MODIFICATION_INSERTED
    # Detecting deleted text
    elif text.startswith(DELETED_OPENING_TAG) and text.endswith(DELETED_CLOSING_TAG):
        return text[len(DELETED_OPENING_TAG):-len(DELETED_CLOSING_TAG)], MODIFICATION_DELETED
    # Unmodified text
    else:
        return text, None

def reformat_doc(docPath, outputFilename, showDeleted):
    """
    This function receives a reference docx file, an output file name,
    and reformats all the specified words.
    """
    # Creating reference and writable documents
    referenceDoc = Document(docPath)
    doc2 = Document()

    # Iterating through each paragraph of the reference document
    for paragraph in referenceDoc.paragraphs:
        # Creating a new paragraph and applying style in the new document
        # based on the reference one
        newParagraph = doc2.add_paragraph()
        newParagraph.alignment = paragraph.alignment

        # Iterating through each run in the current paragraph
        for runXML in paragraph._p.xpath('w:r | w:ins/w:r | w:del/w:r'):
            # Obtaining original run
            run = Run(runXML, paragraph)

            # Getting real text with modification type
            text, modification = get_text_with_type(full_text(runXML))

            # Adding new run with the styles for inserted
            if modification == MODIFICATION_INSERTED:
                add_run_with_style(newParagraph, text, run, highlight=True)
            # Adding new run with the styles for deleted
            elif modification == MODIFICATION_DELETED:
                if showDeleted:
                    add_run_with_style(newParagraph, text, run, color=RED_COLOR, cross=True)
            # Adding new run with the same style as current run
            else:
                add_run_with_style(newParagraph, text, run)

    # Saving the output in the defined document
    doc2.save(outputFilename)

# Main function
if __name__ == '__main__':
    # Param variable
    SHOW_DELETED_PARAM = "--show-deleted"

    # Checking for --help option
    if len(sys.argv) == 2 and sys.argv[1] == "--help":
        print(f"Usage: python docx_diff_reformat.py DOCX_FILE_INPUT DOCX_FILE_OUTPUT [{SHOW_DELETED_PARAM}]")
        print("\tDOCX_FILE_INPUT: path to the .docx file to read from")
        print("\tDOCX_FILE_OUTPUT: path to the .docx file to write to")
        print(f"\t{SHOW_DELETED_PARAM}: optional argument to show deleted content")
        print(f"Example: python docx_diff_reformat.py compare-A-B.docx output.docx {SHOW_DELETED_PARAM}")
        sys.exit(0)

    # Checking parameter count
    if len(sys.argv) < 3:
        print("Error: missing arguments. Please provide both an input and an output .docx files as the first and second arguments.\nUse --help for more information.")
        sys.exit(1)

    # Checking file extension
    docxFile = sys.argv[1]
    for docxFile in [sys.argv[1], sys.argv[2]]:
        if not docxFile.endswith(".docx"):
            print(f"Error: {docxFile} is not a Microsoft Word file. File extension must be .docx.\nUse --help for more information.")
            sys.exit(1)

    # Checking for optional --show-deleted param
    showDeleted = False
    if len(sys.argv) > 3:
        if sys.argv[3] == SHOW_DELETED_PARAM:
            showDeleted = True
        else:
            print(f"Error: '{sys.argv[3]}' param not recognized.\nUse --help for more information.")
            sys.exit(1)
    
    # Reformat Word document
    reformat_doc(sys.argv[1], sys.argv[2], showDeleted)